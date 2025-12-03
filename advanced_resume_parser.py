# -------------------
# Imports & Config
# -------------------
import os
import re
import json
import spacy
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
import logging
from datetime import datetime
import dateutil.parser as dparser
from typing import Dict, List, Any, Tuple
import hashlib
from functools import lru_cache
from rapidfuzz import fuzz, process as rf_process
import numpy as np
from collections import Counter

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.error("spaCy model not found. Please install: python -m spacy download en_core_web_sm")
    nlp = None

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = './temp'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_FILE_SIZE = 5 * 1024 * 1024
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

_PARSE_CACHE: dict[str, dict] = {}


def _hash_file(path: str) -> str:
    hasher = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


class AdvancedResumeParser:
    def __init__(self):
        # Expanded set of non-cities including more tech terms to prevent misclassification
        self.non_cities = set([
            "india", "maharashtra", "uttar pradesh", "madhya pradesh", "karnataka", "tamil nadu", "gujarat", "rajasthan", "punjab", "haryana", "bihar", "west bengal", "andhra pradesh", "telangana", "kerala", "odisha", "jharkhand", "chhattisgarh", "assam", "goa", "manipur", "meghalaya", "mizoram", "nagaland", "sikkim", "tripura", "arunachal pradesh", "jammu", "kashmir", "himachal pradesh", "uttarakhand", "chandigarh", "ladakh", "puducherry", "andaman", "nicobar", "lakshadweep", "daman", "diu", "dadra", "nagar haveli",
            "usa", "united states", "california", "texas", "florida", "new york", "illinois", "ohio", "georgia", "north carolina", "michigan", "new jersey", "virginia", "washington", "arizona", "massachusetts", "tennessee", "indiana", "missouri", "maryland", "wisconsin", "colorado", "minnesota", "south carolina", "alabama", "louisiana", "kentucky", "oregon", "oklahoma", "connecticut", "iowa", "mississippi", "arkansas", "utah", "nevada", "kansas", "new mexico", "nebraska", "west virginia", "idaho", "hawaii", "maine", "new hampshire", "montana", "rhode island", "delaware", "south dakota", "north dakota", "alaska", "vermont",
            "react", "node", "developer", "engineer", "python", "javascript", "java", "c++", "c#", "mern", "full stack", "state", "country", "remote", "onsite", "hybrid",
            "asp.net", "mongodb", "mysql", "numpy", "pandas", "git", "github", "postman", "rest", "apis", "html", "css", "tailwind", "express", "laravel", "sql", "data science", "programming"
        ])
        self._tech_keywords_min = self.non_cities  # Reuse for filtering
        self._job_titles = self.load_job_titles()

    def allowed_file(self, filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX files"""
        try:
            ext = Path(file_path).suffix.lower()
            if ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif ext == '.txt':
                return Path(file_path).read_text(encoding='utf-8', errors='ignore')
            else:
                raise ValueError("Unsupported file format")
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with best-effort strategy: pdfplumber first, then PyPDF2."""
        # Try pdfplumber first for better layout/text capture
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = [page.extract_text(x_tolerance=1, y_tolerance=1) or "" for page in pdf.pages]
                text = "\n".join(pages_text)
                if text and len(text.strip()) >= 50:
                    return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # Fallback: PyPDF2
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"PyPDF2 extraction error: {e}")
            raise
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX with better accuracy"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise
        return text

    def extract_annotations_from_pdf(self, file_path: str) -> List[str]:
        """Extract hyperlink URIs from PDF annotations using pdfplumber."""
        uris = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    for annot in page.annots:
                        if annot.get('subtype') == 'Link':
                            action = annot.get('A')
                            if action and 'URI' in action:
                                uris.append(action['URI'])
        except Exception as e:
            logger.warning(f"Failed to extract annotations: {e}")
        return uris

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information with high accuracy and filter out tech keywords from location"""
        contact = {"email": "", "phone": "", "location": ""}
        # Email extraction
        email_patterns = [
            r'(?i)(?<![\w.-])([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})(?![\w.-])',
        ]
        for pattern in email_patterns:
            matches = re.findall(pattern, text)
            if matches:
                contact["email"] = matches[0]
                break
        # Phone extraction
        phone_patterns = [
            r'(?:\+?91[-\s]?)?[6-9]\d{9}',  # India 10-digit starting 6-9
            r'(?:\+?1[-\s]?)?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',  # US
            r'(?:\+?\d{1,3}[-\s]?)?\d{6,14}',  # E.164 generic
        ]
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                phone = re.sub(r'[^\d+]', '', matches[0])
                if not phone.startswith('+'):
                    if len(phone) == 10 and phone.startswith(('6', '7', '8', '9')):
                        phone = '+91' + phone
                    elif len(phone) == 11 and phone.startswith('0'):
                        phone = '+91' + phone[1:]
                    else:
                        phone = '+' + phone
                contact["phone"] = phone
                break
        # Location extraction using NER, filter out tech keywords (minimal set)
        if nlp:
            top_text = '\n'.join(text.split('\n')[:15])  # Limit to top for contact
            doc = nlp(top_text)
            locations = [ent.text for ent in doc.ents if ent.label_ == "GPE" and ent.text.lower() not in self.non_cities]
            if locations:
                contact["location"] = locations[0]
        return contact

    def extract_city(self, text: str) -> str:
        """Extract city using spaCy NER, limited to top text, filter out states, tech terms, company names, and URLs."""
        if nlp:
            top_text = '\n'.join(text.split('\n')[:15])  # Limit to top lines for contact info
            doc = nlp(top_text)
            for ent in doc.ents:
                if ent.label_ == "GPE":
                    city = ent.text.strip()
                    city_lower = city.lower()
                    # Filter out non-cities, URLs, social/site keywords, states, and empty
                    if (
                        city_lower not in self.non_cities
                        and not re.match(r"https?://", city_lower)
                        and not re.search(r"\b(state|province|region|pradesh|india|united states|usa|gujarat|maharashtra|tamil nadu|karnataka)\b", city_lower)
                        and not re.search(r"\b(github|linkedin|twitter|x\.com|gmail|email)\b", city_lower)
                        and len(city) > 2 and len(city) < 40
                        and not re.search(r"\d", city)
                    ):
                        return city
        return ""

    def extract_city_from_contact_line(self, text: str) -> str:
        """Heuristic for patterns like 'Surat, Gujarat' in top lines: prefer the first token before comma."""
        lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
        states = {"gujarat", "maharashtra", "tamil nadu", "karnataka", "uttar pradesh", "madhya pradesh", "rajasthan", "punjab", "haryana", "bihar", "west bengal"}
        for ln in lines[:10]:
            # If contact line has separators like Github | Linkedin | Surat, Gujarat
            segments = [seg.strip() for seg in re.split(r"\|", ln) if seg.strip()]
            cand = segments[-1] if segments else ln
            if "," in cand:
                parts = [p.strip() for p in cand.split(",")]
                if len(parts) >= 2 and parts[1].lower() in states and 2 <= len(parts[0]) < 30:
                    city = parts[0]
                    if not re.search(r"\b(github|linkedin|twitter|x\.com|gmail|email)\b", city.lower()):
                        return city
        return ""

    def extract_social_links(self, text: str, file_path: str = None) -> dict:
        """Extract LinkedIn, GitHub, Twitter, and portfolio URLs from resume text and PDF annotations."""
        links = {"linkedin": "", "github": "", "twitter": "", "portfolio": ""}
        patterns = {
            "linkedin": r"https?://(www\.)?linkedin\.com/(in|pub)/[A-Za-z0-9\-_/]+/?",
            "github": r"https?://(www\.)?github\.com/[A-Za-z0-9\-_/]+/?",
            "twitter": r"https?://(www\.)?(x|twitter)\.com/[A-Za-z0-9_]+/?",
            "portfolio": r"https?://(?!www\.(linkedin|github|twitter|x)\.com)[\w.-]+\.[a-z]{2,}(/[\w\-./?%&=]*)?"
        }
        # First, extract from text
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                links[key] = match.group(0)

        # If PDF, extract from annotations and classify if not already set
        if file_path and Path(file_path).suffix.lower() == '.pdf':
            ann_uris = self.extract_annotations_from_pdf(file_path)
            for uri in ann_uris:
                uri_lower = uri.lower()
                if "linkedin.com" in uri_lower and not links["linkedin"]:
                    links["linkedin"] = uri
                elif "github.com" in uri_lower and not links["github"]:
                    links["github"] = uri
                elif ("twitter.com" in uri_lower or "x.com" in uri_lower) and not links["twitter"]:
                    links["twitter"] = uri
                elif not any(dom in uri_lower for dom in ["linkedin.com", "github.com", "twitter.com", "x.com"]) and not links["portfolio"]:
                    links["portfolio"] = uri

        # Avoid portfolio being set to LinkedIn/GitHub/Twitter
        for key in ["linkedin", "github", "twitter"]:
            if links[key] and links["portfolio"] == links[key]:
                links["portfolio"] = ""
        return links

    def extract_profile_summary(self, text: str) -> str:
        """Extract a profile summary or objective section if present, else fallback to first paragraph. Never return a URL or empty."""
        summary_patterns = [
            r'(?i)(summary|profile|professional summary|career objective|objective)[:\s]*([\s\S]{30,600}?)(?=\n\n|\n\w+:|$)'
        ]
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                summary = match.group(2).strip()
                summary = re.sub(r'\n+', ' ', summary)
                if summary and not re.match(r'https?://', summary):
                    return summary
        # Fallback: first paragraph (not a section header, not contact info, not a URL)
        paras = [p.strip() for p in text.split('\n\n') if p.strip()]
        for para in paras:
            if len(para) > 30 and not any(x in para.lower() for x in ["email", "phone", "contact", "linkedin", "github", "twitter", "curriculum", "resume", "address"]) and not re.match(r'https?://', para):
                return para.replace('\n', ' ')
        return ""

    def extract_name(self, text: str) -> str:
        """Extract name using spaCy NER and fallback logic"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        # Use spaCy NER for PERSON entities in the first 10 lines
        if nlp:
            for line in lines[:10]:
                doc = nlp(line)
                for ent in doc.ents:
                    if ent.label_ == "PERSON" and 2 <= len(ent.text.split()) <= 4:
                        return ent.text
        # Fallback: first line that looks like a name
        for line in lines[:5]:
            if (
                not any(x in line.lower() for x in ["email", "phone", "contact", "curriculum", "resume"])
                and 2 <= len(line.split()) <= 4
                and line[0].isupper()
            ):
                return line
        return ""

    def _skill_confidence(self, candidate: str, source: str) -> float:
        # High confidence from explicit section, medium from taxonomy/context
        base = 0.9 if source == 'section' else 0.7
        return base if candidate and len(candidate) >= 2 else 0.0

    @lru_cache(maxsize=1)
    def _skills_taxonomy(self) -> List[str]:
        # Try loading from an optional file; otherwise, fall back to a curated list
        candidates = [
            os.path.join(os.path.dirname(__file__), "skills.txt"),
            os.path.join(os.getcwd(), "skills.txt"),
        ]
        for p in candidates:
            try:
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    words = [w.strip() for w in f if w.strip()]
                    if words:
                        return words
            except Exception:
                pass
        return [
            # Programming languages
            "JavaScript","TypeScript","Python","Java","C","C++","C#","Go","Rust","Ruby","PHP","Swift","Kotlin","Scala","R","Matlab",
            # Frontend
            "React","Redux","Next.js","Angular","Vue","Svelte","HTML","CSS","Tailwind","Bootstrap","Sass","Less",
            # Backend
            "Node.js","Express","Django","Flask","Spring","Spring Boot",".NET","ASP.NET","Laravel","Ruby on Rails","GraphQL","REST",
            # Databases
            "MongoDB","MySQL","PostgreSQL","SQLite","Redis","Elasticsearch","DynamoDB","Oracle","SQL Server","NoSQL","SQL",
            # DevOps / Infra
            "Docker","Kubernetes","K8s","Terraform","Ansible","AWS","Azure","GCP","CI/CD","Git","GitHub","GitLab","Jenkins","Nginx","Apache",
            # Data / AI
            "Pandas","NumPy","SciPy","scikit-learn","PyTorch","TensorFlow","Keras","XGBoost","LightGBM","Hugging Face","NLP","Computer Vision","LLM",
            # Mobile / Desktop
            "React Native","Flutter","Android","iOS","Electron",
            # Testing / Quality
            "Jest","Mocha","Chai","Cypress","Playwright","Selenium","Postman",
            # Misc
            "Microservices","Event-Driven","Kafka","RabbitMQ","gRPC","OpenAPI","Swagger","Figma","UX","UI","Agile","Scrum"
        ]

    def extract_skills(self, text: str) -> List[Dict[str, Any]]:
        """Extract skills from explicit section (high confidence) and taxonomy matches (medium)."""
        results: List[Dict[str, Any]] = []
        seen = set()

        # 1) Explicit Skills section
        skills_section = re.search(r'(?i)(skills|technical skills|key skills|technologies)[:\s]*(.*?)(?=\n\n|\n[A-Z]{3,}\b|$)', text, re.DOTALL)
        if skills_section:
            skills_text = skills_section.group(2)
            lines = [l.strip() for l in skills_text.split('\n') if l.strip()]
            for line in lines:
                parts = []
                if ':' in line:
                    _, skills_str = line.split(':', 1)
                    parts = [s.strip() for s in skills_str.split(',') if s.strip()]
                else:
                    parts = [s.strip() for s in re.split(r'[ ,\u2022\-;\/|]+', line) if s.strip()]
                for item in parts:
                    if len(item) < 2 or re.match(r'https?://', item):
                        continue
                    normalized = re.sub(r'\s+', ' ', item)
                    key = normalized.lower()
                    if key in seen:
                        continue
                    seen.add(key)
                    conf = self._skill_confidence(normalized, 'section')
                    if conf >= 0.85:
                        results.append({"name": normalized, "confidence": round(conf, 2)})

        # 2) Taxonomy scan (contextual)
        taxonomy = self._skills_taxonomy()
        text_lower = text.lower()
        for skill in taxonomy:
            pattern = r'(?:^|[^a-z0-9])' + re.escape(skill.lower()) + r'(?:[^a-z0-9]|$)'
            if re.search(pattern, text_lower):
                key = skill.lower()
                if key in seen:
                    continue
                seen.add(key)
                results.append({"name": skill, "confidence": round(self._skill_confidence(skill, 'taxonomy'), 2)})

        return results

    def load_job_titles(self) -> List[str]:
        candidates = [
            os.path.join(os.path.dirname(__file__), "..", "job-titles.txt"),
            os.path.join(os.path.dirname(__file__), "job-titles.txt"),
            os.path.join(os.getcwd(), "job-titles.txt"),
        ]
        for path in candidates:
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    titles = [line.strip() for line in f if line.strip()]
                    if titles:
                        return titles
            except Exception:
                continue
        # Fallback common titles
        return [
            "Software Engineer","Senior Software Engineer","Full Stack Developer","Frontend Developer","Backend Developer",
            "Data Scientist","Machine Learning Engineer","DevOps Engineer","Cloud Engineer","Mobile Developer","Android Developer","iOS Developer",
            "UI/UX Designer","Product Manager","Project Manager","QA Engineer","SDET","Data Engineer","Solutions Architect","Technical Lead"
        ]

    def extract_job_title(self, text: str) -> Tuple[str, float]:
        titles = self._job_titles
        if not titles:
            return "", 0.0
        lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
        header = ' '.join(lines[:15])[:2000]
        # Pull likely phrases
        phrases = set()
        for m in re.finditer(r"[A-Z][A-Za-z0-9\-/+&() ]{2,60}", header):
            phrase = m.group(0).strip()
            if 1 <= len(phrase.split()) <= 8:
                phrases.add(phrase)
        pool = list(phrases) if phrases else [header]
        best_match = rf_process.extractOne(' '.join(pool)[:2000], titles, scorer=fuzz.token_set_ratio)
        if best_match:
            title, score = best_match[0], best_match[1] / 100.0
            if score >= 0.85:
                return title, round(score, 2)
        return "", 0.0

    def _normalize_date_to_yyyy_mm(self, date_str: str) -> str:
        """Convert various date formats to YYYY-MM format required by frontend validation."""
        if not date_str:
            return ""
        date_str = date_str.strip()
        # Already in YYYY-MM format
        if re.match(r'^\d{4}-\d{2}$', date_str):
            return date_str
        # Present/Current
        if date_str.lower() in ['present', 'current']:
            return 'Present'
        # Try parsing common formats: "Jan 2020", "January 2020", "01/2020", "2020"
        try:
            # Handle "Jan 2020" or "January 2020"
            month_year_match = re.match(r'(?i)^([a-z]+)\s+(\d{4})$', date_str)
            if month_year_match:
                month_str, year = month_year_match.groups()
                try:
                    dt = dparser.parse(f"{month_str} 1, {year}", fuzzy=False)
                    return dt.strftime('%Y-%m')
                except:
                    pass
            # Handle "01/2020" or "2020/01"
            slash_match = re.match(r'^(\d{1,2})/(\d{4})$', date_str)
            if slash_match:
                month, year = slash_match.groups()
                return f"{year}-{int(month):02d}"
            # Handle just year "2020"
            if re.match(r'^\d{4}$', date_str):
                return f"{date_str}-01"  # Default to January
        except:
            pass
        return ""

    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience or projects with positions, companies, dates, and descriptions.
        Format dates as YYYY-MM and ensure all required fields are present."""
        experience: List[Dict[str, str]] = []
        # Section detection, including Projects as fallback
        exp_sections = []
        patterns = [
            r'(?i)^(?:work\s+)?experience\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
,            r'(?i)^employment\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
,            r'(?i)^professional\s+experience\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
,            r'(?i)^projects\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
            if match:
                exp_sections.append(match.group(1))
        if not exp_sections:
            exp_sections = [text]

        # Regex for standard experience entry
        entry_regex = re.compile(
            r'(?i)^(?P<position>[A-Za-z .,&\-/+()]+?)\s+(?:at\s+)?(?P<company>[A-Za-z .,&\-/+()]+?)\s+(?P<start>\w+\s+\d{4}|\d{2}/\d{4}|\d{4})\s*(?:[-–—]\s*(?P<end>\w+\s+\d{4}|\d{2}/\d{4}|\d{4}|present|current))?\s*$',
            re.MULTILINE
        )
        # Regex for date lines (e.g., for projects)
        date_regex = re.compile(r'(?i)^(?P<start>\w{3,9}\s+\d{4})\s*-\s*(?P<end>\w{3,9}\s+\d{4}|Present|Current)\s*$')

        for exp_text in exp_sections:
            lines = [ln.strip() for ln in exp_text.split('\n') if ln.strip()]
            i = 0
            while i < len(lines):
                line = lines[i]
                m = entry_regex.match(line)
                if m:
                    raw_start = m.group('start').strip()
                    raw_end = (m.group('end') or 'Present').strip()
                    entry = {
                        "company": m.group('company').strip(),
                        "position": m.group('position').strip(),
                        "start": self._normalize_date_to_yyyy_mm(raw_start),
                        "end": self._normalize_date_to_yyyy_mm(raw_end),
                        "description": "",
                    }
                    # Collect descriptions
                    desc_lines = []
                    j = i + 1
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if entry_regex.match(next_line) or date_regex.match(next_line):
                            break
                        if re.match(r'^\s*[\-•\*]\s+', next_line) or len(next_line) >= 10:
                            desc_lines.append(next_line)
                        j += 1
                    entry["description"] = " ".join(desc_lines)
                    # Only add if required fields are present
                    if entry["company"] and entry["position"] and entry["start"]:
                        experience.append(entry)
                    i = j
                elif date_regex.match(line):
                    # Project-style: date on one line, project name on next
                    dm = date_regex.match(line)
                    raw_start = dm.group('start').strip()
                    raw_end = dm.group('end').strip()
                    entry = {
                        "company": "Project",  # Default company for projects
                        "position": "",
                        "start": self._normalize_date_to_yyyy_mm(raw_start),
                        "end": self._normalize_date_to_yyyy_mm(raw_end),
                        "description": "",
                    }
                    j = i + 1
                    if j < len(lines):
                        next_line = lines[j].strip()
                        if not date_regex.match(next_line) and not entry_regex.match(next_line) and len(next_line) > 5:
                            entry["position"] = next_line  # Project name as position
                            j += 1
                    # Collect descriptions
                    desc_lines = []
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if entry_regex.match(next_line) or date_regex.match(next_line):
                            break
                        if re.match(r'^\s*[\-•\*]\s+', next_line) or len(next_line) >= 10:
                            desc_lines.append(next_line)
                        j += 1
                    entry["description"] = " ".join(desc_lines)
                    # Only add if we have position and valid dates
                    if entry["position"] and entry["start"]:
                        experience.append(entry)
                    i = j
                elif len(line) > 5 and len(line) < 50 and not re.search(r'\d', line) and not re.match(r'^\s*[\-•\*]\s+', line):
                    # Project-style: title on one line, date on next
                    j = i + 1
                    if j < len(lines) and date_regex.match(lines[j]):
                        dm = date_regex.match(lines[j])
                        raw_start = dm.group('start').strip()
                        raw_end = dm.group('end').strip()
                        entry = {
                            "company": "Project",
                            "position": line,
                            "start": self._normalize_date_to_yyyy_mm(raw_start),
                            "end": self._normalize_date_to_yyyy_mm(raw_end),
                            "description": "",
                        }
                        j += 1
                        desc_lines = []
                        while j < len(lines):
                            next_line = lines[j].strip()
                            if entry_regex.match(next_line) or date_regex.match(next_line) or (len(next_line) > 5 and len(next_line) < 50 and not re.search(r'\d', next_line) and not re.match(r'^\s*[\-•\*]\s+', next_line)):
                                break
                            if re.match(r'^\s*[\-•\*]\s+', next_line) or len(next_line) >= 10:
                                desc_lines.append(next_line)
                            j += 1
                        entry["description"] = " ".join(desc_lines)
                        # Only add if valid dates
                        if entry["start"]:
                            experience.append(entry)
                        i = j
                    else:
                        i += 1
                else:
                    i += 1
        return experience

    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education details. Shape entries to {institution, degree, year} for frontend compatibility."""
        education = []
        # Section detection
        edu_sections = []
        patterns = [
            r'(?i)^education\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
,            r'(?i)^academic\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
,            r'(?i)^qualifications?\b[:\s]*(.*)(?=\n\s*\n|\n[A-Z]{3,}\b|$)'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
            if match:
                edu_sections.append(match.group(1))
        if not edu_sections:
            edu_sections = [text]

        # Regexes
        date_range_regex = re.compile(r'(?i)^(?P<start>\w{3,9}\s+\d{4})\s*-\s*(?P<end>\w{3,9}\s+\d{4}|Present)\s*$')
        single_year_regex = re.compile(r'(?i)\b(19|20)\d{2}\b')
        degree_line_regex = re.compile(r'(?i)^(?P<degree>(?:B\.?E\.?|B\.?Tech|B\.?Sc|M\.?Sc|M\.?Tech|Bachelors?|Masters?|MBA|MCA|Ph\.?D\.?|Diploma)[^\n]{0,80})$')

        for edu_text in edu_sections:
            lines = [ln.strip() for ln in edu_text.split('\n') if ln.strip()]
            i = 0
            while i < len(lines):
                line = lines[i]
                entry = {"institution": "", "degree": "", "year": ""}

                # Pattern A: Degree line, then institution, then dates or year
                degm = degree_line_regex.match(line)
                if degm:
                    entry["degree"] = degm.group('degree').strip()
                    j = i + 1
                    if j < len(lines) and not degree_line_regex.match(lines[j]):
                        inst_line = lines[j].strip()
                        if len(inst_line) > 3:
                            entry["institution"] = inst_line
                            j += 1
                    # Look for year on same or next lines
                    year = None
                    scan_k = j
                    while scan_k < min(j + 3, len(lines)) and not year:
                        rngm = date_range_regex.match(lines[scan_k])
                        if rngm:
                            year = re.findall(single_year_regex, rngm.group('end'))
                            year = rngm.group('end')[-4:] if 'Present' not in rngm.group('end') else rngm.group('start')[-4:]
                            break
                        year_match = single_year_regex.search(lines[scan_k])
                        if year_match:
                            year = year_match.group(0)
                            break
                        scan_k += 1
                    if year:
                        entry["year"] = str(year)
                    if entry["degree"] and entry["institution"]:
                        education.append(entry)
                        i = max(j, scan_k)
                        continue

                # Pattern B: Institution first, then degree, then year
                if len(line) > 3 and not any(h in line.lower() for h in ["education", "qualification", "academic"]):
                    # Check if next line is degree
                    j = i + 1
                    degm2 = None
                    if j < len(lines):
                        degm2 = degree_line_regex.match(lines[j])
                    if degm2:
                        entry["institution"] = line
                        entry["degree"] = degm2.group('degree').strip()
                        # Year on next lines
                        k = j + 1
                        year = None
                        while k < min(j + 4, len(lines)) and not year:
                            rngm = date_range_regex.match(lines[k])
                            if rngm:
                                year = rngm.group('end')[-4:] if 'Present' not in rngm.group('end') else rngm.group('start')[-4:]
                                break
                            year_match = single_year_regex.search(lines[k])
                            if year_match:
                                year = year_match.group(0)
                                break
                            k += 1
                        if year:
                            entry["year"] = str(year)
                        education.append(entry)
                        i = k
                        continue

                i += 1
        return education

    def get_tokens(self, text: str) -> List[str]:
        """Preprocess text to get lemmatized tokens without stops."""
        if not nlp:
            return []
        doc = nlp(text.lower())
        tokens = [token.lemma_ for token in doc if token.is_alpha and not token.is_stop]
        return tokens

    def compute_tf_idf(self, corpus: List[str]) -> List[Dict[str, float]]:
        """Compute TF-IDF for a list of documents (space-separated tokens)."""
        n_docs = len(corpus)
        all_words = set()
        doc_words = [doc.split() for doc in corpus]
        for words in doc_words:
            all_words.update(words)
        all_words = list(all_words)
        idf = {}
        for word in all_words:
            df = sum(1 for words in doc_words if word in words)
            idf[word] = np.log(n_docs / (df + 1)) + 1  # Smoothed IDF
        all_words_tf_idf = []
        for doc in doc_words:
            tf = Counter(doc)
            total_terms = len(doc)
            tfidf = {word: (tf[word] / total_terms) * idf[word] for word in tf}
            all_words_tf_idf.append(tfidf)
        return all_words_tf_idf

    def cosine_similarity(self, vec1: Dict[str, float], vec2: Dict[str, float]) -> float:
        """Compute cosine similarity between two TF-IDF vectors."""
        words = set(vec1).union(vec2)
        v1 = np.array([vec1.get(word, 0) for word in words])
        v2 = np.array([vec2.get(word, 0) for word in words])
        dot = np.dot(v1, v2)
        norm1 = np.linalg.norm(v1)
        norm2 = np.linalg.norm(v2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return dot / (norm1 * norm2)

# Instantiate parser after class definition
parser = AdvancedResumeParser()

# Flask API endpoint for resume parsing
@app.route("/parse", methods=["POST"])
def parse_resume_api():
    # Enforce file presence and size
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    cl = request.content_length or 0
    if cl and cl > MAX_FILE_SIZE + 1024:  # small buffer
        return jsonify({"error": "File too large (max 5MB)"}), 413

    if not parser.allowed_file(file.filename):
        return jsonify({"error": "Only PDF, DOCX, or TXT files are supported"}), 400
    job_desc = request.form.get("job_desc", "")  # Optional job description for matching
    filename = secure_filename(file.filename)
    temp_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(temp_path)
    try:
        file_hash = _hash_file(temp_path)
        cached = _PARSE_CACHE.get(file_hash)
        if cached:
            os.remove(temp_path)
            return jsonify(cached), 200

        text = parser.extract_text(temp_path)
        name = parser.extract_name(text)
        job_title, job_title_conf = parser.extract_job_title(text)
        contact = parser.extract_contact_info(text)
        city = parser.extract_city(text) or parser.extract_city_from_contact_line(text)
        social_links = parser.extract_social_links(text, temp_path)
        skills = parser.extract_skills(text)
        experience = parser.extract_experience(text)
        education = parser.extract_education(text)
        profile_summary = parser.extract_profile_summary(text)

        # Simple confidences for key fields
        def conf_from_presence(value: str) -> float:
            return 0.9 if value and len(value) >= 3 else 0.0

        # Fallback name from email if necessary
        if (not name) and contact.get("email"):
            local = contact.get("email").split("@")[0]
            local = re.sub(r'[._-]+', ' ', local)
            name = ' '.join(w.capitalize() for w in local.split() if w)

        result = {
            "name": name,
            "name_confidence": conf_from_presence(name),
            "title": job_title,
            "title_confidence": job_title_conf,
            "email": contact.get("email", ""),
            "email_confidence": conf_from_presence(contact.get("email", "")),
            "phone": contact.get("phone", ""),
            "phone_confidence": conf_from_presence(contact.get("phone", "")),
            "location": city or contact.get("location", ""),
            "location_confidence": conf_from_presence(city or contact.get("location", "")),
            "skills": skills,  # list of {name, confidence}
            "experience": experience,
            "education": education,
            "about": profile_summary,
            "about_confidence": conf_from_presence(profile_summary),
            "linkedin": social_links.get("linkedin", ""),
            "github": social_links.get("github", ""),
            "twitter": social_links.get("twitter", ""),
            "portfolio": social_links.get("portfolio", ""),
        }

        # Add ML-based match score if job_desc provided
        if job_desc and nlp:
            resume_text = ' '.join([profile_summary] + [s['name'] for s in skills] + [e['description'] for e in experience if e['description']])
            resume_tokens = parser.get_tokens(resume_text)
            job_tokens = parser.get_tokens(job_desc)
            if resume_tokens and job_tokens:
                corpus = [' '.join(job_tokens), ' '.join(resume_tokens)]
                tfidfs = parser.compute_tf_idf(corpus)
                score = parser.cosine_similarity(tfidfs[0], tfidfs[1])
                result["match_score"] = round(score, 2)
                # Arbitrary confidence based on score
                result["match_confidence"] = 0.9 if score > 0.5 else 0.5

        # Cache result by file hash in-memory
        _PARSE_CACHE[file_hash] = result
        os.remove(temp_path)
        return jsonify(result), 200
    except Exception as e:
        logger.exception(f"Parse failed: {e}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({"error": str(e)}), 500

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "model": "en_core_web_sm", "cache_entries": len(_PARSE_CACHE)})

@app.route("/cache/clear", methods=["POST"])
def clear_cache():
    _PARSE_CACHE.clear()
    return jsonify({"cleared": True})

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5001))
    app.run(host="0.0.0.0", port=port)